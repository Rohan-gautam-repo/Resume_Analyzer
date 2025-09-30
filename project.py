import streamlit as st
import os
import google.generativeai as genai
import pandas as pd
import docx
from docx import Document
from docx.shared import Pt, RGBColor
import pdfplumber
import spacy
import io
import re
import time
from datetime import datetime
from dotenv import load_dotenv

# ✅ Ensure Streamlit Page Config is FIRST
st.set_page_config(
    page_title="AI Resume Analyzer", 
    layout="wide", 
    page_icon="📄",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better UI
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem !important;
        color: #1E88E5;
        text-align: center;
        margin-bottom: 1rem;
    }
    .subheader {
        font-size: 1.5rem;
        color: #0D47A1;
        margin-top: 2rem;
        border-bottom: 2px solid #1E88E5;
        padding-bottom: 0.5rem;
    }
    .card {
        background-color: #f8f9fa;
        border-radius: 15px;
        padding: 25px;
        box-shadow: 0 6px 12px rgba(0,0,0,0.1);
        margin-bottom: 25px;
        border: 1px solid #e0e0e0;
    }
    .highlight {
        color: #1E88E5;
        font-weight: bold;
    }
    .progress-bar {
        height: 20px;
        border-radius: 10px;
    }
    .skill-tag {
        background: linear-gradient(135deg, #1E88E5, #0D47A1);
        color: white;
        padding: 10px 15px;
        margin: 6px;
        border-radius: 20px;
        text-align: center;
        font-weight: 500;
        font-size: 0.9rem;
        display: inline-block;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        transition: transform 0.2s;
    }
    .skill-tag:hover {
        transform: translateY(-2px);
    }
    .template-preview {
        border: 2px solid #1E88E5;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
    .comparison-container {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        border-radius: 15px;
        padding: 20px;
        margin: 20px 0;
    }
    .ai-suggestions {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-radius: 15px;
        padding: 25px;
        margin: 20px 0;
    }
    .success-message {
        background: linear-gradient(135deg, #4CAF50, #45a049);
        color: white;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        margin: 15px 0;
    }
    .feature-box {
        background: white;
        border-radius: 12px;
        padding: 20px;
        margin: 10px 0;
        border-left: 4px solid #1E88E5;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 20px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 60px;
        padding: 0 20px;
        border-radius: 10px 10px 0 0;
        background: linear-gradient(135deg, #1E88E5, #0D47A1);
        color: white;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: linear-gradient(135deg, #4CAF50, #45a049);
    }
</style>
""", unsafe_allow_html=True)

# ✅ Load API Key from .env
load_dotenv()
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")

# ✅ Verify API Key
if not GEMINI_API_KEY:
    st.error("🚨 API Key is missing! Ensure GOOGLE_API_KEY is set in the .env file.")
    st.stop()

# ✅ Configure Google Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# 📌 Function to get available models
@st.cache_data(ttl=3600)  # Cache for 1 hour
def get_available_models():
    try:
        models = genai.list_models()
        available_models = []
        for model in models:
            if 'generateContent' in model.supported_generation_methods:
                available_models.append(model.name)
        return available_models
    except Exception as e:
        st.warning(f"Could not fetch available models: {str(e)}")
        return []

# ✅ Load NLP Model
@st.cache_resource
def load_nlp_model():
    try:
        return spacy.load("en_core_web_sm")
    except Exception as e:
        st.warning("⚠️ Installing SpaCy model 'en_core_web_sm'... This may take a moment.")
        try:
            import subprocess
            import sys
            # Try to download the model
            subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
            return spacy.load("en_core_web_sm")
        except Exception as install_error:
            st.error("❌ Failed to install SpaCy model. Using basic text processing instead.")
            return None

nlp = load_nlp_model()

# 📌 Function to Extract Text from PDF
def extract_text_from_pdf(pdf_file):
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = '\n'.join([page.extract_text() for page in pdf.pages if page.extract_text()])
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

# 📌 Function to Extract Text from DOCX
def extract_text_from_docx(docx_file):
    try:
        doc = docx.Document(docx_file)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

# 📌 Function to Parse Resume
def parse_resume(uploaded_file):
    file_extension = uploaded_file.name.split(".")[-1].lower()
    if file_extension == "pdf":
        return extract_text_from_pdf(uploaded_file), file_extension
    elif file_extension == "docx":
        return extract_text_from_docx(uploaded_file), file_extension
    else:
        return None, None

# 📌 Function to Extract Skills using NLP
def extract_skills(text):
    skills = set()
    
    # Enhanced common skills list
    common_skills = [
        "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "swift", "kotlin", "go", "rust",
        "sql", "mysql", "postgresql", "mongodb", "oracle", "nosql", "firebase",
        "machine learning", "deep learning", "natural language processing", "nlp", "computer vision",
        "react", "angular", "vue", "node.js", "express", "django", "flask", "spring", "laravel",
        "docker", "kubernetes", "aws", "azure", "gcp", "cloud computing", "devops", "cicd",
        "linux", "unix", "bash", "powershell", "git", "github", "gitlab", "bitbucket",
        "excel", "power bi", "tableau", "data analysis", "data visualization", "statistics", "r",
        "leadership", "communication", "project management", "agile", "scrum", "kanban",
        "html", "css", "sass", "less", "responsive design", "ui/ux", "figma", "sketch",
        "tensorflow", "pytorch", "keras", "pandas", "numpy", "scikit-learn", "matplotlib",
        "rest api", "graphql", "oauth", "authentication", "blockchain", "cybersecurity"
    ]
    
    text_lower = text.lower()
    
    # If nlp model is available, use it for better token extraction
    if nlp:
        try:
            doc = nlp(text_lower)
            # Extract single word skills
            for token in doc:
                if token.text in common_skills:
                    skills.add(token.text)
        except:
            pass  # Fall back to basic text matching
    
    # Extract multi-word skills (always do this regardless of nlp availability)
    for skill in common_skills:
        if skill.lower() in text_lower:
            skills.add(skill)
    
    return list(skills)

# 📌 Function to Get AI-Powered Resume Suggestions from Gemini API
def ai_resume_improvement_gemini(resume_text, job_description=None):
    try:
        # Get available models
        available_models = get_available_models()
        
        # Try different model names in order of preference
        preferred_models = [
            "gemini-1.5-flash",
            "gemini-1.0-pro", 
            "gemini-pro",
            "models/gemini-1.5-flash",
            "models/gemini-1.0-pro",
            "models/gemini-pro"
        ]
        
        # Find the first available preferred model
        model_to_use = None
        for preferred in preferred_models:
            for available in available_models:
                if preferred in available or available.endswith(preferred):
                    model_to_use = available
                    break
            if model_to_use:
                break
        
        # If no preferred model found, use the first available one
        if not model_to_use and available_models:
            model_to_use = available_models[0]
        
        if not model_to_use:
            return "Error: No suitable Gemini model available. Please check your API key or try again later."
        
        # Create the model
        model = genai.GenerativeModel(model_to_use)
        
        # Enhanced prompt with job matching if available
        if job_description:
            prompt = f"""
            You are a professional resume consultant. Analyze this resume and provide specific improvements to make it more effective and ATS-friendly. 
            
            RESUME:
            {resume_text}
            
            JOB DESCRIPTION:
            {job_description}
            
            Please provide your analysis in the following format:
            
            ## Overall Assessment
            [Provide a brief overall assessment]
            
            ## Strengths
            - [Strength 1]
            - [Strength 2]
            - [Strength 3]
            
            ## Areas for Improvement
            - [Area 1]
            - [Area 2]
            - [Area 3]
            
            ## Specific Suggestions to Match Job Description
            [Detailed suggestions to better align with the job]
            
            ## Improved Resume
            [Provide a complete, improved version of the resume]
            """
        else:
            prompt = f"""
            You are a professional resume consultant. Analyze this resume and provide specific improvements to make it more effective and ATS-friendly.
            
            RESUME:
            {resume_text}
            
            Please provide your analysis in the following format:
            
            ## Overall Assessment
            [Provide a brief overall assessment]
            
            ## Strengths
            - [Strength 1]
            - [Strength 2]
            - [Strength 3]
            
            ## Areas for Improvement
            - [Area 1]
            - [Area 2]
            - [Area 3]
            
            ## Improved Resume
            [Provide a complete, improved version of the resume]
            """
        
        response = model.generate_content(prompt)
        
        if hasattr(response, 'text'):
            return response.text
        else:
            return "No suggestions available."
    
    except Exception as e:
        return f"Error calling Gemini API: {str(e)}"

# 📌 Function to Generate ATS Score
def calculate_ats_score(resume_text):
    # Initialize base score
    score = 100
    feedback = []
    
    # Check for essential sections (more comprehensive)
    essential_sections = {
        'experience': ['experience', 'work experience', 'employment', 'career', 'professional experience'],
        'education': ['education', 'academic', 'degree', 'university', 'college'],
        'skills': ['skills', 'technical skills', 'competencies', 'expertise'],
        'summary': ['summary', 'objective', 'profile', 'about']
    }
    
    missing_sections = []
    text_lower = resume_text.lower()
    
    for section_name, keywords in essential_sections.items():
        found = any(keyword in text_lower for keyword in keywords)
        if not found:
            missing_sections.append(section_name)
            score -= 15  # Increased penalty for missing sections
            feedback.append(f"Missing {section_name} section")
    
    # Check for proper formatting (enhanced)
    formatting_checks = {
        'bullet_points': {
            'check': len(re.findall(r'[•·▪▫◦‣⁃]|\u2022|\u2023|\d+\.|\*\s', resume_text)) >= 5,
            'penalty': 8,
            'message': 'Insufficient bullet points for better readability'
        },
        'contact_info': {
            'check': bool(re.search(r'\b[\w\.-]+@[\w\.-]+\.\w{2,4}\b', resume_text)),
            'penalty': 12,
            'message': 'Missing email address'
        },
        'phone_number': {
            'check': bool(re.search(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', resume_text)),
            'penalty': 8,
            'message': 'Missing phone number'
        },
        'length': {
            'check': 300 < len(resume_text) < 2000,
            'penalty': 10,
            'message': 'Resume length not optimal (too short or too long)'
        },
        'professional_keywords': {
            'check': len([word for word in ['achieved', 'managed', 'developed', 'created', 'implemented', 'improved', 'increased', 'reduced', 'led', 'coordinated'] if word in text_lower]) >= 3,
            'penalty': 6,
            'message': 'Lacks strong action verbs and professional keywords'
        },
        'quantifiable_achievements': {
            'check': len(re.findall(r'\d+%|\d+\+|\d+k|\d+m|\$\d+|\d+ years', resume_text)) >= 2,
            'penalty': 8,
            'message': 'Missing quantifiable achievements and metrics'
        }
    }
    
    formatting_issues = []
    for check_name, check_data in formatting_checks.items():
        if not check_data['check']:
            score -= check_data['penalty']
            formatting_issues.append(check_name)
            feedback.append(check_data['message'])
    
    # Enhanced skills assessment
    skills = extract_skills(resume_text)
    if len(skills) < 8:
        penalty = (8 - len(skills)) * 3
        score -= penalty
        feedback.append(f"Limited technical skills identified ({len(skills)} found, recommend 8+)")
    
    # Check for ATS-friendly formatting
    ats_unfriendly_elements = {
        'tables': len(re.findall(r'\|.*\|', resume_text)),
        'graphics_references': len(re.findall(r'image|figure|chart|graph', text_lower)),
        'special_characters': len(re.findall(r'[^\w\s\-.,;:()\n\r@#%&+=/]', resume_text))
    }
    
    if ats_unfriendly_elements['tables'] > 2:
        score -= 5
        feedback.append("Contains tables that may not be ATS-friendly")
    
    if ats_unfriendly_elements['graphics_references'] > 0:
        score -= 3
        feedback.append("References to graphics/images that ATS cannot read")
    
    if ats_unfriendly_elements['special_characters'] > 10:
        score -= 4
        feedback.append("Contains special characters that may cause ATS parsing issues")
    
    # Ensure score stays within 0-100
    score = max(0, min(100, score))
    
    return {
        "score": score,
        "missing_sections": missing_sections,
        "formatting_issues": formatting_issues,
        "skills_found": skills,
        "feedback": feedback,
        "recommendations": generate_ats_recommendations(score, missing_sections, formatting_issues, len(skills))
    }

# 📌 Function to generate ATS recommendations
def generate_ats_recommendations(score, missing_sections, formatting_issues, skills_count):
    recommendations = []
    
    if score >= 80:
        recommendations.append("✅ Excellent ATS compatibility! Your resume should pass most ATS systems.")
    elif score >= 60:
        recommendations.append("⚠️ Good ATS compatibility with room for improvement.")
    else:
        recommendations.append("❌ Significant improvements needed for ATS compatibility.")
    
    if missing_sections:
        recommendations.append(f"📝 Add missing sections: {', '.join(missing_sections)}")
    
    if skills_count < 8:
        recommendations.append("🔧 Include more relevant technical skills and keywords")
    
    if 'bullet_points' in formatting_issues:
        recommendations.append("📌 Use more bullet points to improve readability")
    
    if 'contact_info' in formatting_issues:
        recommendations.append("📧 Ensure contact information is clearly visible")
    
    recommendations.append("🎯 Use standard section headings like 'Work Experience', 'Education', 'Skills'")
    recommendations.append("📊 Include quantifiable achievements with specific numbers and percentages")
    
    return recommendations

# 📌 Function to extract the improved resume section from AI suggestions
def extract_improved_resume(ai_suggestions):
    if "## Improved Resume" in ai_suggestions:
        parts = ai_suggestions.split("## Improved Resume")
        if len(parts) > 1:
            improved_text = parts[1].strip()
            # Clean up markdown formatting
            improved_text = clean_resume_text(improved_text)
            return improved_text
    return None

# 📌 Function to clean up resume text formatting
def clean_resume_text(text):
    """Remove markdown formatting and improve text presentation"""
    # Remove ** bold formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Remove * italic formatting
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Remove ### headings and convert to proper format
    text = re.sub(r'###\s*(.*?)\n', r'\1\n', text)
    text = re.sub(r'##\s*(.*?)\n', r'\1\n', text)
    text = re.sub(r'#\s*(.*?)\n', r'\1\n', text)
    
    # Clean up bullet points
    text = re.sub(r'^-\s*', '• ', text, flags=re.MULTILINE)
    
    # Remove extra whitespace and newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    
    return text

# 📌 Function to convert text to DOCX with template styling
def text_to_docx(text, template_name="Modern Professional"):
    doc = Document()
    
    # Apply template styles
    if template_name == "Classic Corporate":
        # Classic template styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(12)
        heading2_style.font.bold = True
        heading2_style.font.italic = False
    elif template_name == "Modern Professional":
        # Modern template styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(16)
        heading_style.font.color.rgb = RGBColor(0x1E, 0x88, 0xE5)  # Blue
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Calibri'
        heading2_style.font.size = Pt(13)
        heading2_style.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)  # Darker blue
        heading2_style.paragraph_format.space_after = Pt(4)
    elif template_name == "Creative Designer":
        # Creative template styles
        style = doc.styles['Normal']
        style.font.name = 'Georgia'
        style.font.size = Pt(11)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Georgia'
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0x8E, 0x24, 0xAA)  # Purple
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Georgia'
        heading2_style.font.size = Pt(14)
        heading2_style.font.color.rgb = RGBColor(0x6A, 0x1B, 0x99)  # Dark purple
    elif template_name == "Tech Specialist":
        # Tech template styles
        style = doc.styles['Normal']
        style.font.name = 'Consolas'
        style.font.size = Pt(10)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0x00, 0x7A, 0xCC)  # Tech blue
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Arial'
        heading2_style.font.size = Pt(12)
        heading2_style.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)  # Darker tech blue
    elif template_name == "Executive Leader":
        # Executive template styles
        style = doc.styles['Normal']
        style.font.name = 'Garamond'
        style.font.size = Pt(12)
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Garamond'
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark executive blue
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Garamond'
        heading2_style.font.size = Pt(14)
        heading2_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)  # Medium executive blue
    elif template_name == "Minimalist Clean":
        # Minimalist template styles
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.15
        # Heading 1
        heading_style = doc.styles['Heading 1']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(6)
        # Heading 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = 'Arial'
        heading2_style.font.size = Pt(12)
        heading2_style.font.bold = False
        heading2_style.font.italic = True
    
    # Process text content
    paragraphs = text.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            if para.strip().startswith('#'):
                level = len(re.match(r'^#+', para.strip()).group(0))
                heading_text = para.strip().lstrip('#').strip()
                doc.add_heading(heading_text, level=level if level <= 9 else 1)
            else:
                lines = para.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith(('- ', '* ', '• ')):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.\s', line):
                        doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')
                    else:
                        doc.add_paragraph(line)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ✅ Main App UI
st.markdown("<h1 class='main-header'>📄 AI-Powered Resume Analyzer</h1>", unsafe_allow_html=True)

# Add a welcome message and instructions
st.markdown("""
<div style='text-align: center; margin-bottom: 2rem;'>
    <p style='font-size: 1.2rem; color: #666;'>AI-powered resume analysis with instant ATS scoring, skills extraction, and professional improvement suggestions</p>
</div>
""", unsafe_allow_html=True)

# Create tabs for different sections
tab1, tab2, tab3 = st.tabs(["📤 Upload & Analyze", "📊 ATS Score", "✨ AI Improvements"])

# Check if we should auto-switch to AI Improvements tab
if st.session_state.get("switch_to_ai_tab", False):
    st.session_state.switch_to_ai_tab = False
    # Use JavaScript to switch to the AI Improvements tab
    st.markdown("""
    <script>
        setTimeout(function() {
            const tabs = document.querySelectorAll('[data-baseweb="tab"]');
            if (tabs.length >= 3) {
                tabs[2].click();
            }
        }, 100);
    </script>
    """, unsafe_allow_html=True)

with tab1:
    # 📂 Resume Upload with file size limit
    uploaded_file = st.file_uploader(
        "📂 Upload Your Resume (PDF/DOCX - Max 10MB)", 
        type=["pdf", "docx"],
        help="Upload your resume in PDF or DOCX format. Maximum file size: 10MB"
    )
    
    if uploaded_file:
        # Check file size (10MB limit)
        if uploaded_file.size > 10 * 1024 * 1024:  # 10MB in bytes
            st.error("❌ File size too large! Please upload a file smaller than 10MB.")
            st.stop()
        # Store the original file
        if "original_file" not in st.session_state:
            file_bytes = uploaded_file.getvalue()
            st.session_state.original_file = io.BytesIO(file_bytes)
            st.session_state.original_filename = uploaded_file.name
        
        # Extract Resume Text
        if "resume_text" not in st.session_state or st.session_state.get("resume_filename") != uploaded_file.name:
            with st.spinner("📄 Extracting resume content..."):
                resume_text, file_extension = parse_resume(uploaded_file)
                st.session_state.resume_text = resume_text
                st.session_state.file_extension = file_extension
                st.session_state.resume_filename = uploaded_file.name
                
                # Extract Skills
                st.session_state.skills = extract_skills(resume_text) if resume_text else []
                
                # Reset AI suggestions when a new file is uploaded
                if "ai_suggestions" in st.session_state:
                    del st.session_state.ai_suggestions
                if "improved_resume" in st.session_state:
                    del st.session_state.improved_resume
                
            st.success("✅ Resume uploaded successfully!")
        
        # Display extracted text
        st.markdown("<h2 class='subheader'>📄 Resume Content</h2>", unsafe_allow_html=True)
        st.text_area("Resume Text", st.session_state.resume_text, height=250)
        
        # Display skills
        skills = st.session_state.skills
        if skills:
            st.markdown("<h2 class='subheader'>🔍 Extracted Skills</h2>", unsafe_allow_html=True)
            
            # Create a more visual representation of skills with darker color
            cols = st.columns(3)
            for i, skill in enumerate(skills):
                col_idx = i % 3
                cols[col_idx].markdown(f"<div class='skill-tag'>{skill}</div>", unsafe_allow_html=True)
        
        # AI Analysis button
        if st.button("🚀 Analyze Resume with AI", type="primary", use_container_width=True):
            if not st.session_state.resume_text or len(st.session_state.resume_text.strip()) < 100:
                st.error("⚠️ Resume content seems too short. Please upload a complete resume for better analysis.")
            else:
                with st.spinner("🤖 AI is analyzing your resume... This may take a moment..."):
                    # Add a small delay to make the spinner visible
                    time.sleep(1)
                    
                    # Get AI suggestions
                    ai_suggestions = ai_resume_improvement_gemini(st.session_state.resume_text)
                    
                    if ai_suggestions.startswith("Error"):
                        st.error(f"❌ {ai_suggestions}")
                        st.info("💡 Try refreshing the page or check your internet connection.")
                    else:
                        st.session_state.ai_suggestions = ai_suggestions
                        
                        # Extract the improved resume part
                        improved_resume = extract_improved_resume(ai_suggestions)
                        if improved_resume:
                            st.session_state.improved_resume = improved_resume
                        
                        # Set flag to auto-switch to AI Improvements tab
                        st.session_state.switch_to_ai_tab = True
                        st.success("✅ Analysis complete! Switching to AI Improvements tab...")
                        st.balloons()  # Celebratory animation
                        st.rerun()  # Refresh to trigger tab switch

with tab2:
    # ATS-specific file uploader with size limit
    ats_file = st.file_uploader(
        "📄 Upload Resume for ATS Analysis (Max 10MB)", 
        type=["pdf", "docx"],
        key="ats_upload",
        help="Upload your resume for ATS compatibility analysis. Maximum file size: 10MB"
    )
    
    if ats_file:
        # Check file size (10MB limit)
        if ats_file.size > 10 * 1024 * 1024:  # 10MB in bytes
            st.error("❌ File size too large! Please upload a file smaller than 10MB.")
            st.stop()
        with st.spinner("🔍 Analyzing ATS Compatibility..."):
            # Parse resume
            resume_text, _ = parse_resume(ats_file)
            
            if resume_text:
                # Calculate ATS score
                ats_result = calculate_ats_score(resume_text)
                st.session_state.ats_result = ats_result
                
                # Display results
                score = ats_result["score"]
                
                st.markdown("<h2 class='subheader'>📊 ATS Compatibility Score</h2>", unsafe_allow_html=True)
                
                # Visual score meter
                col1, col2 = st.columns([1, 3])
                with col1:
                    st.markdown(f"""
                    <div style="text-align: center;">
                        <div style="font-size: 3rem; font-weight: bold; color: {'#4CAF50' if score >= 70 else '#FFC107' if score >= 50 else '#F44336'};">
                            {score}%
                        </div>
                        <div style="font-size: 0.8rem; color: gray;">
                            ATS Readiness Score
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                with col2:
                    st.markdown(f"""
                    <div style="margin-top: 20px;">
                        <div style="background-color: #e0e0e0; border-radius: 10px; height: 20px;">
                            <div style="background-color: {'#4CAF50' if score >= 70 else '#FFC107' if score >= 50 else '#F44336'}; 
                                        width: {score}%; 
                                        height: 100%; 
                                        border-radius: 10px;">
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Score interpretation
                    if score >= 70:
                        st.markdown("✅ **Excellent!** Your resume is ATS-friendly.")
                    elif score >= 50:
                        st.markdown("⚠️ **Moderate.** Some improvements needed for better ATS performance.")
                    else:
                        st.markdown("❌ **Needs Work.** Significant improvements required for ATS systems.")
                
                # Detailed analysis
                st.markdown("#### 🔍 Detailed Analysis")
                
                # Feedback
                if ats_result.get("feedback"):
                    st.markdown("**Issues Found:**")
                    for feedback_item in ats_result["feedback"]:
                        st.markdown(f"- {feedback_item}")
                
                # Recommendations
                if ats_result.get("recommendations"):
                    st.markdown("#### 💡 Recommendations")
                    for rec in ats_result["recommendations"]:
                        st.markdown(f"- {rec}")
                
                # Missing sections
                if ats_result["missing_sections"]:
                    st.markdown(f"**Missing Sections:** {', '.join(ats_result['missing_sections'])}")
                
                # Skills found
                st.markdown("#### ✅ Identified Skills")
                if ats_result["skills_found"]:
                    cols = st.columns(3)
                    for i, skill in enumerate(ats_result["skills_found"]):
                        cols[i%3].markdown(f"<div class='skill-tag'>{skill}</div>", unsafe_allow_html=True)
                else:
                    st.info("No technical skills detected. Consider adding relevant skills to your resume.")
                
            else:
                st.error("Error processing uploaded file")

with tab3:
    
    if uploaded_file and "resume_text" in st.session_state:
        if "ai_suggestions" in st.session_state:
            # Display AI suggestions in a beautiful format
            st.markdown("<h2 class='subheader'>✨ AI-Powered Improvement Suggestions</h2>", unsafe_allow_html=True)
            
            # Create a styled container for AI suggestions
            st.markdown("<div class='ai-suggestions'>", unsafe_allow_html=True)
            st.markdown("### 🤖 AI Analysis Results")
            st.markdown(st.session_state.ai_suggestions)
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Check if we have improved resume
            if "improved_resume" in st.session_state and st.session_state.improved_resume:
                st.markdown("<h2 class='subheader'>📝 Accept AI Improvements</h2>", unsafe_allow_html=True)
                
                # Create three columns for better layout
                col1, col2, col3 = st.columns([1, 2, 1])
                
                with col2:
                    # Template Selection with better styling
                    st.markdown("#### 🎨 Choose Your Resume Template")
                    selected_template = st.selectbox(
                        "",
                        ["Modern Professional", "Classic Corporate", "Creative Designer", "Tech Specialist", "Executive Leader", "Minimalist Clean"],
                        index=0,
                        key="template_select",
                        help="Select a professional template for your improved resume"
                    )
                    
                    # Template descriptions in feature boxes
                    template_descriptions = {
                        "Modern Professional": ("🎯", "Contemporary design with clean lines and professional blue accents. Perfect for tech, consulting, and modern industries."),
                        "Classic Corporate": ("📋", "Traditional format with Times New Roman font. Ideal for conservative industries like finance, law, and government."),
                        "Creative Designer": ("�", "Stylish layout with creative elements. Great for design, marketing, and creative professionals."),
                        "Tech Specialist": ("�", "Tech-focused format with modern typography. Optimized for software engineers and IT professionals."),
                        "Executive Leader": ("👔", "Distinguished design for senior leadership roles. Perfect for C-level executives and management positions."),
                        "Minimalist Clean": ("✨", "Simple and elegant with focus on content clarity. Suitable for any industry that values clean presentation.")
                    }
                    
                    icon, desc = template_descriptions[selected_template]
                    st.markdown(f"""
                    <div class='feature-box'>
                        <h4>{icon} {selected_template} Template</h4>
                        <p>{desc}</p>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Before/After comparison with better styling
                st.markdown("<h3 style='text-align: center; margin: 30px 0;'>📊 Before vs After Comparison</h3>", unsafe_allow_html=True)
                
                st.markdown("<div class='comparison-container'>", unsafe_allow_html=True)
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("##### 📄 Original Resume")
                    st.text_area(
                        "", 
                        st.session_state.resume_text, 
                        height=250, 
                        disabled=True, 
                        key="original_display",
                        help="Your original resume content"
                    )
                
                with col2:
                    st.markdown("##### ✨ AI-Improved Resume")
                    st.text_area(
                        "", 
                        st.session_state.improved_resume, 
                        height=250, 
                        disabled=True, 
                        key="improved_display",
                        help="AI-enhanced version of your resume"
                    )
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Action buttons with better styling
                st.markdown("<h3 style='text-align: center; margin: 30px 0;'>🚀 Download Your Improved Resume</h3>", unsafe_allow_html=True)
                
                col1, col2, col3 = st.columns([1, 2, 1])
                with col2:
                    if st.button("✅ Generate & Download Improved Resume", type="primary", use_container_width=True):
                        with st.spinner("🔄 Creating your professional resume..."):
                            # Convert improved resume to DOCX
                            docx_buffer = text_to_docx(st.session_state.improved_resume, selected_template)
                            
                            # Get original filename without extension
                            filename = os.path.splitext(st.session_state.original_filename)[0]
                            current_date = datetime.now().strftime("%Y-%m-%d")
                            new_filename = f"{filename}_AI_improved_{current_date}_{selected_template.lower()}.docx"
                            
                            st.markdown("<div class='success-message'>", unsafe_allow_html=True)
                            st.markdown("🎉 **Success!** Your improved resume is ready for download!")
                            st.markdown("</div>", unsafe_allow_html=True)
                            
                            # Offer for download
                            st.download_button(
                                label="📥 Download Improved Resume (DOCX)",
                                data=docx_buffer,
                                file_name=new_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                
                # Original resume download option
                st.markdown("---")
                st.markdown("#### 📎 Original Resume Download")
                if st.session_state.file_extension:
                    mime_types = {
                        "pdf": "application/pdf",
                        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    }
                    
                    mime_type = mime_types.get(st.session_state.file_extension, "application/octet-stream")
                    
                    col1, col2, col3 = st.columns([1, 2, 1])
                    with col2:
                        st.download_button(
                            label=f"📥 Download Original Resume ({st.session_state.file_extension.upper()})",
                            data=st.session_state.original_file,
                            file_name=st.session_state.original_filename,
                            mime=mime_type,
                            use_container_width=True
                        )
            else:
                # If no improved resume available
                st.markdown("""
                <div class='feature-box'>
                    <h4>🔍 Analysis Complete</h4>
                    <p>AI has analyzed your resume and provided suggestions above. The improved version will appear here once generated.</p>
                </div>
                """, unsafe_allow_html=True)
        else:
            # No AI suggestions yet
            st.markdown("""
            <div class='feature-box'>
                <h3>🚀 Ready for AI Analysis</h3>
                <p>Upload your resume in the <strong>Upload & Analyze</strong> tab and click <strong>'Analyze Resume with AI'</strong> to get intelligent improvement suggestions powered by Google Gemini.</p>
                <hr>
                <h4>What you'll get:</h4>
                <ul>
                    <li>📊 Comprehensive resume analysis</li>
                    <li>✨ AI-powered improvement suggestions</li>
                    <li>📝 Professional resume rewrite</li>
                    <li>🎨 Multiple template options</li>
                    <li>📥 Instant DOCX download</li>
                </ul>
            </div>
            """, unsafe_allow_html=True)
    else:
        # No resume uploaded
        st.markdown("""
        <div class='feature-box'>
            <h3>📄 Upload Your Resume First</h3>
            <p>To get AI-powered improvement suggestions, please upload your resume (PDF or DOCX format) in the <strong>Upload & Analyze</strong> tab.</p>
            <hr>
            <h4>🎯 Our AI Resume Analyzer provides:</h4>
            <ul>
                <li><strong>Smart Analysis:</strong> Identifies strengths and areas for improvement</li>
                <li><strong>ATS Optimization:</strong> Ensures your resume passes Applicant Tracking Systems</li>
                <li><strong>Professional Rewriting:</strong> Creates an enhanced version of your resume</li>
                <li><strong>Template Selection:</strong> Choose from multiple professional formats</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

# Footer with better styling
st.markdown("---")
st.markdown("""
<div style='text-align: center; margin-top: 3rem; padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 15px;'>
    <div style='color: white; font-size: 1.1rem; margin-bottom: 1rem;'>
        <strong>🎯 AI Resume Analyzer</strong>
    </div>
    <div style='color: #f0f0f0; font-size: 0.9rem; margin-bottom: 1rem;'>
        Powered by Google Gemini AI • Built with Streamlit
    </div>
    <div style='color: #d0d0d0; font-size: 0.8rem;'>
        © 2025 AI Resume Analyzer | Helping professionals succeed with AI-powered resume optimization
    </div>
</div>
""", unsafe_allow_html=True)
